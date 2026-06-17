"""
DOCX extraction (blueprint §6) — python-docx.

Walks body content in true document order (paragraphs *and* tables
interleaved, which ``document.paragraphs`` alone doesn't give you), turning
each paragraph into a ``TextBlock``.

Signals python-docx gives us that PDF doesn't:

* ``paragraph_style`` (e.g. "Heading 1", "Normal", "List Bullet") — the
  single strongest heading signal for DOCX.
* run-level bold/italic and explicit font size (``None`` when inherited).
* numbering properties (``numPr``) and ``left_indent`` for list/indent.

Tables: a 2-column table that looks like a *layout container* (sidebar
template) is split into two columns; every other table is flattened
row-major into the current column with ``is_table_cell=True``.
``page_number`` is always 0.
"""

from __future__ import annotations

from collections import defaultdict
from typing import Optional

from docx import Document
from docx.document import Document as _Document
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from ..intermediate_representation import ExtractedDocument, TextBlock
from ..normalizer import strip_leading_bullet
from .base import INDENT_UNIT_EMU, MAX_INDENT_LEVEL


def _iter_block_items(parent):
    """Yield Paragraph and Table children of *parent* in document order."""
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError(f"unsupported parent: {type(parent)!r}")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


class DocxExtractor:
    def extract(self, filepath: str) -> ExtractedDocument:
        document = Document(filepath)
        self._blocks: list[TextBlock] = []
        self._order: dict[int, int] = defaultdict(int)
        self._column_count = 1

        for item in _iter_block_items(document):
            if isinstance(item, Paragraph):
                self._emit_paragraph(item, column_index=0)
            elif isinstance(item, Table):
                self._handle_table(item)

        return ExtractedDocument(
            source_format="docx",
            blocks=self._blocks,
            column_count=self._column_count,
        )

    # -- paragraphs ------------------------------------------------------
    def _emit_paragraph(
        self, para: Paragraph, column_index: int, is_table_cell: bool = False
    ) -> None:
        raw = para.text
        if not raw or not raw.strip():
            return  # spacing paragraph

        text, bulleted = strip_leading_bullet(raw.strip())
        text = text.strip()
        if not text:
            return

        style_name = para.style.name if para.style is not None else None
        is_list = bulleted or self._is_list_paragraph(para, style_name)
        bold, italic = self._run_emphasis(para)
        size = self._font_size(para)
        level = self._indent_level(para, is_list)

        order = self._order[column_index]
        self._order[column_index] = order + 1
        self._blocks.append(
            TextBlock(
                text=text,
                order_index=order,
                column_index=column_index,
                page_number=0,
                font_size=size,
                paragraph_style=style_name,
                is_bold=bold,
                is_italic=italic,
                is_list_item=is_list,
                indentation_level=level,
                is_table_cell=is_table_cell,
            )
        )

    @staticmethod
    def _is_list_paragraph(para: Paragraph, style_name: Optional[str]) -> bool:
        if style_name and "list" in style_name.lower():
            return True
        pPr = para._p.pPr
        if pPr is not None and pPr.numPr is not None:
            return True
        return False

    @staticmethod
    def _run_emphasis(para: Paragraph) -> tuple[bool, bool]:
        """Bold/italic by character-weighted majority of the paragraph's runs."""
        runs = [r for r in para.runs if r.text and r.text.strip()]
        if not runs:
            return False, False
        total = sum(len(r.text) for r in runs)
        bold = sum(len(r.text) for r in runs if r.bold)
        italic = sum(len(r.text) for r in runs if r.italic)
        return (bold * 2 >= total, italic * 2 >= total)

    @staticmethod
    def _font_size(para: Paragraph) -> Optional[float]:
        for run in para.runs:
            if run.text and run.text.strip() and run.font.size is not None:
                return float(run.font.size.pt)
        return None

    @staticmethod
    def _indent_level(para: Paragraph, is_list: bool) -> int:
        left = para.paragraph_format.left_indent
        if left is not None:
            try:
                return max(0, min(MAX_INDENT_LEVEL, int(int(left) / INDENT_UNIT_EMU)))
            except (TypeError, ValueError):
                pass
        # Fall back to the list nesting level if present.
        pPr = para._p.pPr
        if pPr is not None and pPr.numPr is not None and pPr.numPr.ilvl is not None:
            ilvl = pPr.numPr.ilvl.get(qn("w:val"))
            if ilvl is not None:
                try:
                    return max(0, min(MAX_INDENT_LEVEL, int(ilvl)))
                except ValueError:
                    return 0
        return 0

    # -- tables ----------------------------------------------------------
    def _handle_table(self, table: Table) -> None:
        if self._is_layout_table(table):
            self._column_count = max(self._column_count, 2)
            for row in table.rows:
                for col_idx, cell in self._unique_cells(row):
                    self._emit_cell(cell, column_index=min(col_idx, 1))
        else:
            # Content table -> flatten row-major into the current main column.
            for row in table.rows:
                for _, cell in self._unique_cells(row):
                    self._emit_cell(cell, column_index=0, is_table_cell=True)

    def _emit_cell(
        self, cell: _Cell, column_index: int, is_table_cell: bool = False
    ) -> None:
        for item in _iter_block_items(cell):
            if isinstance(item, Paragraph):
                self._emit_paragraph(
                    item, column_index=column_index, is_table_cell=is_table_cell
                )
            elif isinstance(item, Table):
                # Nested table -> always treat as content.
                for row in item.rows:
                    for _, nested in self._unique_cells(row):
                        self._emit_cell(nested, column_index=column_index, is_table_cell=True)

    @staticmethod
    def _unique_cells(row):
        """Yield (column_index, cell) skipping horizontally-merged duplicates."""
        seen = set()
        for idx, cell in enumerate(row.cells):
            key = id(cell._tc)
            if key in seen:
                continue
            seen.add(key)
            yield idx, cell

    @staticmethod
    def _is_layout_table(table: Table) -> bool:
        """Heuristic: a 2-column table acting as a page layout container.

        True when the table has exactly two columns and at least one cell is
        "content rich" (several paragraphs, or a heading-styled paragraph) —
        i.e. a sidebar template rather than a small skills/contact grid.
        Acknowledged as heuristic in the blueprint (§6, §14).
        """
        try:
            ncols = len(table.columns)
        except Exception:
            ncols = len(table.rows[0].cells) if table.rows else 0
        if ncols != 2:
            return False
        for row in table.rows:
            for cell in row.cells:
                paras = [p for p in cell.paragraphs if p.text.strip()]
                if len(paras) >= 3:
                    return True
                for p in paras:
                    name = (p.style.name if p.style is not None else "") or ""
                    low = name.lower()
                    if low.startswith("heading") or low in ("title", "subtitle"):
                        return True
        return False
