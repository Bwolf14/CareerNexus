"""
Generate small, representative sample resumes for integration testing.

We can't ship real resumes, so we synthesize files that exercise the layout
variety the blueprint's integration corpus calls for (§13): single-column
PDF, two-column/sidebar PDF, DOCX using Word heading styles, DOCX using
manual bold formatting only, a sidebar (layout-table) DOCX, and a DOCX with
no detectable headings (the "failed" path).

Run as a script to (re)populate ``tests/sample_resumes/``:

    python -m resume_parser.tests.generate_samples
"""

from __future__ import annotations

from pathlib import Path

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt


# --------------------------------------------------------------------------
# PDF helpers
# --------------------------------------------------------------------------
def _write_pdf(path: str, lines: list[tuple], width: float = 612, height: float = 792) -> None:
    """lines: (x, y, text, size, bold)."""
    doc = fitz.open()
    page = doc.new_page(width=width, height=height)
    for x, y, text, size, bold in lines:
        page.insert_text(
            (x, y), text, fontsize=size, fontname="hebo" if bold else "helv"
        )
    doc.save(path)
    doc.close()


def make_single_column_pdf(path: str) -> None:
    x = 60
    y = 60
    lines: list[tuple] = []

    def add(text, size=11, bold=False, dx=0, gap=18):
        nonlocal y
        lines.append((x + dx, y, text, size, bold))
        y += gap

    add("ALEX SMITH", size=20, bold=True, gap=24)
    add("alex.smith@example.com | (587) 555-0199 | Calgary, AB", size=10, gap=14)
    add("linkedin.com/in/alexsmith | github.com/alexsmith", size=10, gap=24)

    add("SUMMARY", size=14, bold=True, gap=20)
    add("Network technician with three years of helpdesk and infrastructure", size=11, gap=14)
    add("experience supporting small-business clients.", size=11, gap=24)

    add("WORK EXPERIENCE", size=14, bold=True, gap=20)
    add("Network Technician — Foothills IT Inc. (Jan 2022 - Present)", size=11, bold=True, gap=16)
    add("•  Maintained switching and routing for 12 client sites.", size=11, dx=18, gap=16)
    add("•  Cut average ticket resolution time by 30%.", size=11, dx=18, gap=20)
    add("Help Desk Analyst — TechCare Solutions (2020 - 2022)", size=11, bold=True, gap=16)
    add("•  Triaged tier-1 and tier-2 support requests.", size=11, dx=18, gap=24)

    add("EDUCATION", size=14, bold=True, gap=20)
    add("Diploma in Network Administration — SAIT (2018 - 2020), GPA: 3.7/4.0", size=11, bold=True, gap=24)

    add("SKILLS", size=14, bold=True, gap=20)
    add("Networking: TCP/IP, VLANs, BGP, OSPF", size=11, gap=14)
    add("Tools: Wireshark, pfSense, Cisco IOS", size=11, gap=24)

    add("CERTIFICATIONS", size=14, bold=True, gap=20)
    add("CompTIA Network+ — CompTIA (2021)", size=11, gap=24)

    add("AWARDS", size=14, bold=True, gap=20)
    add("Dean's Honour List, 2019", size=11, gap=14)

    _write_pdf(path, lines)


def make_two_column_pdf(path: str) -> None:
    # Right column starts at x=200 with a clear gutter from the narrow left
    # sidebar; lines are kept short enough to stay within the page width
    # (text drawn past the page edge gets clipped on extraction).
    lines: list[tuple] = []

    # Left sidebar column (kept narrow so a clear gutter remains).
    lx, ly = 50, 70
    for text, size, bold, gap in [
        ("CONTACT", 13, True, 18),
        ("jordan@example.com", 9, False, 13),
        ("(403) 555-0143", 9, False, 13),
        ("Lethbridge, AB", 9, False, 24),
        ("SKILLS", 13, True, 18),
        ("Python", 10, False, 13),
        ("Docker", 10, False, 13),
        ("Linux", 10, False, 13),
        ("SQL", 10, False, 13),
    ]:
        lines.append((lx, ly, text, size, bold))
        ly += gap

    # Right main column.
    rx, ry = 200, 70
    for text, size, bold, gap in [
        ("JORDAN LEE", 18, True, 26),
        ("EXPERIENCE", 14, True, 20),
        ("Software Developer · Bow Valley Corp (2021 - Present)", 11, True, 16),
        ("•  Built internal tooling in Python and Go.", 11, False, 16),
        ("•  Owned CI/CD pipelines for three services.", 11, False, 22),
        ("EDUCATION", 14, True, 20),
        ("BSc in Computer Science · University of Lethbridge (2017-2021)", 11, True, 22),
    ]:
        lines.append((rx, ry, text, size, bold))
        ry += gap

    _write_pdf(path, lines)


def make_scanned_like_pdf(path: str) -> None:
    """An (almost) empty PDF — stands in for a scanned/image resume."""
    doc = fitz.open()
    doc.new_page(width=612, height=792)
    doc.save(path)
    doc.close()


# --------------------------------------------------------------------------
# DOCX helpers
# --------------------------------------------------------------------------
def make_heading_styles_docx(path: str) -> None:
    doc = Document()

    title = doc.add_paragraph()
    run = title.add_run("Priya Patel")
    run.bold = True
    run.font.size = Pt(20)

    doc.add_paragraph("priya.patel@example.com | (780) 555-0170 | Edmonton, AB")
    doc.add_paragraph("linkedin.com/in/priyapatel")

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(
        "Detail-oriented data analyst with experience in reporting and ETL "
        "pipelines for healthcare organizations."
    )

    doc.add_heading("Experience", level=1)
    p = doc.add_paragraph()
    p.add_run("Data Analyst — Northern Health Ltd. (2021 - Present)").bold = True
    doc.add_paragraph("Automated weekly reporting, saving 10 hours per week.", style="List Bullet")
    doc.add_paragraph("Built dashboards used by 40+ staff.", style="List Bullet")

    doc.add_heading("Education", level=1)
    p = doc.add_paragraph()
    p.add_run("Bachelor of Science in Statistics — University of Alberta (2017 - 2021)").bold = True

    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Programming: Python, R, SQL")
    doc.add_paragraph("Tools: Power BI, Tableau, Excel")

    doc.add_heading("Languages", level=1)  # not in dictionary -> additional_sections
    doc.add_paragraph("English (native), Hindi (fluent)")

    doc.save(path)


def make_manual_bold_docx(path: str) -> None:
    doc = Document()

    def bold_line(text, size=None, all_caps_heading=False):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        if size:
            r.font.size = Pt(size)
        return p

    bold_line("MORGAN TAYLOR", size=18)
    doc.add_paragraph("morgan.taylor@example.com | (250) 555-0122 | Victoria, BC")

    bold_line("WORK EXPERIENCE")  # ALL CAPS bold -> heading via segmenter rule
    p = doc.add_paragraph()
    p.add_run("Warehouse Supervisor — Pacific Logistics (2019 - 2023)").bold = True
    doc.add_paragraph("Managed a team of 15 across two shifts.")
    doc.add_paragraph("Reduced shipping errors by 22%.")

    bold_line("EDUCATION")
    p = doc.add_paragraph()
    p.add_run("Diploma in Supply Chain Management — Camosun College (2017 - 2019)").bold = True

    bold_line("SKILLS")
    doc.add_paragraph("Inventory Management, Forklift Certified, WHMIS, Scheduling")

    doc.save(path)


def make_sidebar_docx(path: str) -> None:
    """Two-column layout table (sidebar template)."""
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    left, right = table.rows[0].cells

    left.paragraphs[0].add_run("CONTACT").bold = True
    left.add_paragraph("sam@example.com")
    left.add_paragraph("(604) 555-0111")
    left.add_paragraph("Burnaby, BC")
    left.add_paragraph().add_run("SKILLS").bold = True
    left.add_paragraph("JavaScript, React, Node.js, CSS")

    right.paragraphs[0].add_run("SAM RIVERA").bold = True
    p = right.add_paragraph()
    p.add_run("EXPERIENCE").bold = True
    p = right.add_paragraph()
    p.add_run("Front-End Developer — Coastal Web Co. (2020 - Present)").bold = True
    right.add_paragraph("Shipped a component library used across 5 products.")
    p = right.add_paragraph()
    p.add_run("EDUCATION").bold = True
    p = right.add_paragraph()
    p.add_run("BSc in Computer Science — UBC (2016 - 2020)").bold = True

    doc.save(path)


def make_no_headers_docx(path: str) -> None:
    """Plain prose, no styled/bold/large headings -> 'failed' status path."""
    doc = Document()
    doc.add_paragraph("Casey Nguyen")
    doc.add_paragraph(
        "I have worked in customer service and retail for several years and am "
        "looking for new opportunities in the technology sector."
    )
    doc.add_paragraph(
        "Most recently I worked at a coffee shop where I trained new staff and "
        "handled inventory ordering."
    )
    doc.add_paragraph("Contact me at casey.nguyen@example.com or (902) 555-0133.")
    doc.save(path)


_PDF_BUILDERS = {
    "single_column.pdf": make_single_column_pdf,
    "two_column.pdf": make_two_column_pdf,
    "scanned_like.pdf": make_scanned_like_pdf,
}
_DOCX_BUILDERS = {
    "heading_styles.docx": make_heading_styles_docx,
    "manual_bold.docx": make_manual_bold_docx,
    "sidebar.docx": make_sidebar_docx,
    "no_headers.docx": make_no_headers_docx,
}


def generate_all(out_dir: str) -> dict[str, str]:
    """Generate every sample into *out_dir*; return {name: path}."""
    out = Path(out_dir)
    out.mkdir(parents=True, exist_ok=True)
    paths: dict[str, str] = {}
    for name, builder in {**_PDF_BUILDERS, **_DOCX_BUILDERS}.items():
        p = out / name
        builder(str(p))
        paths[name] = str(p)
    return paths


if __name__ == "__main__":
    target = Path(__file__).parent / "sample_resumes"
    made = generate_all(str(target))
    for name in made:
        print("wrote", name)
